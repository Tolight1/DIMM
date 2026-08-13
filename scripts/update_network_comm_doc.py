from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SOURCE = "network_comm_source.docx"
OUTPUT = "network_comm_updated.docx"
TABLE_WIDTH_DXA = 9072


def set_cell_text(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=8.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_width(cell, width_dxa):
    cell.width = Inches(width_dxa / 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def mark_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def shade_header(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "E2F0D9")


def format_table(table, widths, header=True):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid_cols = table._tbl.tblGrid.gridCol_lst
    for grid_col, width in zip(grid_cols, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_borders(cell)
    if header:
        mark_header(table.rows[0])
        for cell in table.rows[0].cells:
            shade_header(cell)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True


def add_caption(document, text):
    paragraph = document.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(text)
    return paragraph


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    for i, value in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], value, WD_ALIGN_PARAGRAPH.CENTER)
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 3 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, alignment)
    format_table(table, widths)
    return table


document = Document(SOURCE)

# Replace the old 32-byte frame description with the implementation currently in src.
for paragraph in document.paragraphs:
    if paragraph.text.startswith("气象站DATA区固定为32字节"):
        paragraph.text = (
            "当前项目气象站监控帧的DATA区固定为52字节（12个big-endian float32和1个big-endian uint32状态位），"
            "因此LEN固定为85字节（0x00000055），整帧固定为93字节。"
            "DST_ADDR固定为01 03 03 02 00 00，SRC_ADDR固定为01 03 03 05 00 00，MSG_TYPE固定为0x07。"
            "请气象站确认以下字段顺序、单位和状态位定义。"
        )

# Update the existing frame table (table 9) in place.
frame_rows = [
    ["0", "SOF", "4", "bytes[4]", "49 96 02 D2", "固定帧头"],
    ["4", "LEN", "4", "uint32", "00 00 00 55", "85字节；从DST_ADDR至EOF，不含SOF和LEN"],
    ["8", "DST_ADDR", "6", "bytes[6]", "01 03 03 02 00 00", "综合控制任务支持端点"],
    ["14", "SRC_ADDR", "6", "bytes[6]", "01 03 03 05 00 00", "DIMM设备监控端点"],
    ["20", "MSG_TYPE", "1", "uint8", "07", "状态主动上报"],
    ["21", "SEQ", "4", "uint32", "X", "循环计数，大端序"],
    ["25", "TIMESTAMP", "8", "uint64", "X", "UTC毫秒时标，表示本次数据采样时刻"],
    ["33", "DATA", "52", "—", "见表10", "12个float32字段 + 1个uint32设备状态"],
    ["85", "CRC", "4", "uint32", "X", "CRC-32/ISO-HDLC，计算LEN至DATA末尾"],
    ["89", "EOF", "4", "bytes[4]", "B6 69 FD 2E", "固定帧尾"],
]
frame_table = document.tables[4]
frame_headers = ["帧偏移", "字段", "长度", "类型", "固定值/取值", "说明"]
for cell, value in zip(frame_table.rows[0].cells, frame_headers):
    set_cell_text(cell, value, WD_ALIGN_PARAGRAPH.CENTER)
    shade_header(cell)
for row, values in zip(frame_table.rows[1:], frame_rows):
    prevent_row_split(row)
    for cell, value in zip(row.cells, values):
        set_cell_text(cell, value, WD_ALIGN_PARAGRAPH.CENTER if cell is not row.cells[-1] else WD_ALIGN_PARAGRAPH.LEFT)

# Append the current DATA definition and status bitmask for station review.
# Use direct formatting because the source document's custom heading style has a vertical layout.
heading = document.add_paragraph()
heading.paragraph_format.page_break_before = True
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
heading_run = heading.add_run("当前项目上报字段（请气象站确认）")
heading_run.bold = True
heading_run.font.name = "黑体"
heading_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
heading_run.font.size = Pt(14)

note = document.add_paragraph()
note.add_run(
    "以下内容对应当前项目 src/CommProtocol.cpp 和 src/DIMM.Results.cpp 的实际发送顺序。"
    "所有数值字段均按大端序传输；float32无效时发送IEEE-754 NaN。deviceStatus为可组合的状态位，"
    "0x00000000表示当前未检测到异常。"
)

add_caption(document, "表 10 当前项目 DATA 区字段定义")
data_rows = [
    ["0", "temperature", "4", "float32", "°C", "温度传感器读数；无效时NaN"],
    ["4", "humidity", "4", "float32", "%RH", "相对湿度；无效时NaN"],
    ["8", "pressure", "4", "float32", "hPa", "气压；无效时NaN"],
    ["12", "r0", "4", "float32", "cm", "大气参数反演结果；无效时NaN"],
    ["16", "seeing", "4", "float32", "arcsec", "视宁度；无效时NaN"],
    ["20", "theta0", "4", "float32", "arcsec", "大气相干角；无效时NaN"],
    ["24", "tau0", "4", "float32", "ms", "大气相干时间；无效、欠分辨或过期时NaN"],
    ["28", "peakBrightnessA", "4", "float32", "DN", "相机A星点峰值亮度；无有效星点时NaN"],
    ["32", "peakBrightnessB", "4", "float32", "DN", "相机B星点峰值亮度；无有效星点时NaN"],
    ["36", "exposureTimeA", "4", "float32", "us", "相机A曝光时间；无效时NaN"],
    ["40", "exposureTimeB", "4", "float32", "us", "相机B曝光时间；无效时NaN"],
    ["44", "frameRate", "4", "float32", "Hz", "两相机共用帧频；连续模式为两相机读回值平均，触发模式为触发器频率；无效时NaN"],
    ["48", "deviceStatus", "4", "uint32", "bitmask", "设备状态位；按大端序传输，见表11"],
]
add_table(document, ["DATA偏移", "字段", "长度", "类型", "单位", "说明"], data_rows, [800, 1900, 600, 1100, 1100, 3572])

add_caption(document, "表 11 deviceStatus 设备状态位定义")
status_rows = [
    ["0x00000001", "相机A连接异常", "相机A未打开或连接丢失"],
    ["0x00000002", "相机B连接异常", "相机B未打开或连接丢失"],
    ["0x00000004", "触发器异常", "触发器响应超时，或已启用但未正常运行"],
    ["0x00000008", "相机A采集异常", "相机A采集超时或帧丢失"],
    ["0x00000010", "相机B采集异常", "相机B采集超时或帧丢失"],
    ["0x00000020", "相机A未找到星", "相机A当前没有有效星点质心"],
    ["0x00000040", "相机B未找到星", "相机B当前没有有效星点质心"],
    ["0x00000080", "相机A星点亮度不足", "相机A星点峰值低于当前低亮度阈值"],
    ["0x00000100", "相机B星点亮度不足", "相机B星点峰值低于当前低亮度阈值"],
    ["0x00000200", "环境传感器异常", "启用环境传感器但未取得有效温湿压数据"],
    ["0x00000400", "曝光设置异常", "曝光时间无效或实际值与请求值偏差过大"],
    ["0x00000800", "帧频异常", "帧频无效，或两相机帧频一致性异常"],
    ["0x00001000", "测量结果异常", "大气参数无效或数据过期"],
    ["0x00002000", "结果保存异常", "本地测量结果文件无法写入"],
]
add_table(document, ["位掩码", "状态名称", "置位条件"], status_rows, [1600, 2200, 5272])

footer_note = document.add_paragraph()
footer_note.paragraph_format.space_before = Pt(0)
footer_note.paragraph_format.space_after = Pt(0)
footer_run = footer_note.add_run(
    "审核重点：请确认93字节整帧、LEN=85、DATA顺序/单位、NaN处理、deviceStatus位定义及CRC范围（LEN至DATA末尾）。"
)
footer_run.font.size = Pt(8.5)

document.save(OUTPUT)
print(OUTPUT)
